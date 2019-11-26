myURL ='https://www.dw.com/de/deutsch-lernen/nachrichten/s-8030'

import pandas as pd 
import numpy as np 
import matplotlib.pyplot as plt
import requests
from bs4 import BeautifulSoup # to extract data from html files
from docx import Document
import re


def save_mp3(audio_link):
    mp3_html=requests.get(audio_link).content
    mp3_soup=BeautifulSoup(mp3_html,'html.parser')
    mp3_name=mp3_soup.h1.text
    mp3_href=mp3_soup.find('a',{'target':'hiddenDownloadIframe'})['href']

    temp_name=mp3_name[-26:-19]
    if(temp_name !='langsam'):
        temp_name='original'

    name_list=re.findall(r'\b\d+\b',mp3_name)   
    mp3_name=name_list[2]+name_list[1]+name_list[0]+" "+temp_name+'.mp3'

    mp3_file =requests.get(mp3_href)
    with open(mp3_name,'wb') as f:
        f.write(mp3_file.content)
    

    

html=requests.get(myURL)
html=html.content
soup = BeautifulSoup(html,'html.parser')


base ='https://www.dw.com/'
all_links = soup.find_all('a',{'class':'overlayLink'})
for link in all_links:
    audio_link=base+link['href']   #get address for nachrichten audio
    save_mp3(audio_link)
print('已保存音频文件')

'''
get today news address link
save dw nachrichten to doc file 
'''
my_div=soup.find_all('div',{'class':'news'})
news_link=base+my_div[1].a['href']
news_html=requests.get(news_link)
news_html=news_html.content
news_soup =BeautifulSoup(news_html,'html.parser')

doc =Document()

news_div = news_soup.find('div',{'class':'longText'})
para_text=""
for news_p_text in news_div.find_all('p'):
    if(news_p_text.strong!=None):       #if have strong text
        para_text=news_p_text.strong.text
        news_para=doc.add_paragraph("")
        news_run=news_para.add_run(para_text)
        news_run.bold=True
    else:
        para_text=news_p_text.text      #if not strong text
        news_para=doc.add_paragraph(para_text)
doc_name=news_link[23:31]
doc_name=doc_name[-4:]+doc_name[2:4]+doc_name[:2]+'.docx'
doc.save(doc_name)
print('已保存word文档')
